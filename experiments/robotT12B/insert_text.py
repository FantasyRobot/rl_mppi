from docx import Document
from docx.shared import RGBColor
import sys
import os

def insert_text_with_color(file_path, output_path):
    try:
        doc = Document(file_path)
        
        # Text to be inserted
        cdf_text_title = "Advanced Distance Fields for Safe MPPI Sampling"
        cdf_text_content = (
            " While traditional MPPI methods rely on Euclidean Signed Distance Fields (SDF) in the workspace for collision checking, "
            "recent research has focused on integrating high-dimensional Configuration-space Distance Fields (CDF) directly into the MPPI cost function to improve sampling efficiency. "
            "Neural SDFs [22, 23] and learned implicit representations [24] have been proposed to enable differentiable distance queries, "
            "facilitating gradient-based trajectory optimization within the MPPI framework. However, integrating these neural fields often incurs high inference costs during the massive parallel sampling of MPPI. "
            "To address this, Kernel-based CDF methods [25] and Riemannian Motion Policies (RMP) [26] have been introduced to define metrics on curved manifolds, "
            "offering more accurate distance estimations for revolute manipulators than flat Euclidean approximations. "
            "Unlike these data-driven approaches that may suffer from generalization errors, our method leverages a geometrically consistent "
            "CDF that directly maps workspace obstacles to C-space constraints. This provides a dense and precise safety gradient that naturally guides MPPI sampling toward collision-free regions "
            "without the need for extensive neural network training or expensive online queries."
        )

        rl_mppi_text_title = "Constrained and RL-Guided MPPI"
        rl_mppi_text_content = (
            " Hybrid architectures combining RL and MPC have become a frontier in safe control, specifically leveraging offline policies to warm-start online MPPI planning. "
            "RL-Driven MPPI [8] and TD-MPC [14] utilize offline policies as the nominal mean for MPPI sampling, significantly reducing the variance and horizon requirements compared to random initialization. "
            "Building on this, recent works have integrated safety constraints more tightly into this hybrid loop. "
            "Constraint-Discounted MPPI (CD-MPPI) [27] introduces a constraint-dependent discount factor to suppress trajectories with high violation risks during the importance sampling phase. "
            "Similarly, Safety Critic architectures [28, 29] train auxiliary networks to estimate the probability of future constraint violations, "
            "which are then used as terminal costs to guide MPPI sampling toward safe long-term states. Residual RL frameworks [30] further enhance this by learning task-specific corrections on top of a nominal MPPI safety controller. "
            "Our PG-MPPI framework distinguishes itself by constructing a multi-level safety system: it uses CD-SAC to internalize safety preferences into the guiding policy, "
            "while explicitly enforcing feasibility through MPPI’s online receding horizon optimization, "
            "effectively bridging the gap between probabilistic learning and deterministic safety requirements."
        )

        # Locate the specific paragraph after which to insert
        target_text = "ignoring the complexity of the topological structure of the joint space."
        
        found = False
        for i, para in enumerate(doc.paragraphs):
            if target_text in para.text:
                found = True
                # Insert a new paragraph after the target paragraph
                new_para = para.insert_paragraph_before("")
                
                # Add Title 1
                run_title1 = new_para.add_run(cdf_text_title + "\n")
                run_title1.bold = True
                run_title1.font.color.rgb = RGBColor(255, 0, 0) # Red color
                
                # Add Content 1
                run_content1 = new_para.add_run(cdf_text_content + "\n\n")
                run_content1.font.color.rgb = RGBColor(255, 0, 0) # Red color
                
                # Add Title 2
                run_title2 = new_para.add_run(rl_mppi_text_title + "\n")
                run_title2.bold = True
                run_title2.font.color.rgb = RGBColor(255, 0, 0) # Red color
                
                # Add Content 2
                run_content2 = new_para.add_run(rl_mppi_text_content)
                run_content2.font.color.rgb = RGBColor(255, 0, 0) # Red color
                
                break
        
        if not found:
            print("Target paragraph not found. Appending to the end of Introduction.")
            # If not found, append to the end of the document (or you might want to force it into Introduction if you can identify it)
            # For robustness, let's try to find "1. Introduction" section and append at the end of it if possible, 
            # but simplified here to just append if specific target fails.
            
            # Trying to find the end of Introduction section loosely by "2. Preliminaries"
            prelim_found = False
            for i, para in enumerate(doc.paragraphs):
                if "2. Preliminaries" in para.text:
                    prelim_found = True
                    prev_para = doc.paragraphs[i-1]
                    new_para = prev_para.insert_paragraph_before("") # Insert before Preliminaries

                    # Same insertion logic
                    run_title1 = new_para.add_run(cdf_text_title + "\n")
                    run_title1.bold = True
                    run_title1.font.color.rgb = RGBColor(255, 0, 0)
                    run_content1 = new_para.add_run(cdf_text_content + "\n\n")
                    run_content1.font.color.rgb = RGBColor(255, 0, 0)
                    run_title2 = new_para.add_run(rl_mppi_text_title + "\n")
                    run_title2.bold = True
                    run_title2.font.color.rgb = RGBColor(255, 0, 0)
                    run_content2 = new_para.add_run(rl_mppi_text_content)
                    run_content2.font.color.rgb = RGBColor(255, 0, 0)
                    break
            
            if not prelim_found:
                 print("Could not find '2. Preliminaries' either. Aborting insertion to avoid messing up document structure.")
                 return

        doc.save(output_path)
        print(f"Successfully modified document saved to: {output_path}")

    except Exception as e:
        print(f"Error modifying file: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python insert_text.py <input_docx> <output_docx>")
    else:
        insert_text_with_color(sys.argv[1], sys.argv[2])
